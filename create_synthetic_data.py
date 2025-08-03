import pandas as pd
import numpy as np
from faker import Faker
from PIL import Image, ImageDraw, ImageFont
import os
import docx
from fpdf import FPDF

# Initialize Faker
fake = Faker()

# Create directories if they don't exist
os.makedirs('data/applicants', exist_ok=True)

def create_mock_emirates_id(applicant_name, file_path):
    # (This function remains unchanged)
    width, height = 400, 250
    img = Image.new('RGB', (width, height), color = 'lightgrey')
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 15)
    except IOError:
        font = ImageFont.load_default()
    d.text((10,10), "United Arab Emirates", fill=(0,0,0), font=font)
    d.text((10,30), "Emirates Identity Card", fill=(0,0,0), font=font)
    d.text((10, 80), f"Name: {applicant_name}", fill=(0,0,0), font=font)
    d.text((10, 100), f"IDN: 784-1990-1234567-1", fill=(0,0,0), font=font)
    d.text((10, 120), f"Nationality: Emirati", fill=(0,0,0), font=font)
    img.save(file_path)

def create_resume_content(name, employment_years):
    return (
        f"CV for {name}\n\n"
        f"Contact: {fake.email()}\n\n"
        "Experience:\n"
        f"- {fake.job()} at {fake.company()} ({employment_years} years)\n"
        "- Previous job at another company.\n\n"
        "Skills:\n- Python\n- Project Management\n- Data Analysis"
    )

def create_docx_resume(content, file_path):
    doc = docx.Document()
    doc.add_heading('Curriculum Vitae', 0)
    doc.add_paragraph(content)
    doc.save(file_path)

def create_pdf_resume(content, file_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Use multi_cell to handle newlines properly
    pdf.multi_cell(0, 10, txt=content)
    pdf.output(file_path)

def create_applicant_data(num_applicants=10):
    """Generates a dataset of applicants and their documents."""
    data = []
    for i in range(num_applicants):
        applicant_id = 1000 + i
        name = fake.name()
        income = np.random.randint(2000, 15000)
        family_size = np.random.randint(1, 7)
        employment_years = np.random.randint(0, 20)
        
        # Determine eligibility for the training data
        if income < 5000 and family_size > 3:
            eligibility = "Approve"
        elif income < 8000 and family_size > 1:
            eligibility = np.random.choice(["Approve", "Decline"], p=[0.6, 0.4])
        else:
            eligibility = "Decline"

        applicant_folder = f'data/applicants/{applicant_id}'
        os.makedirs(applicant_folder, exist_ok=True)
        
        # --- Create Documents ---
        # 1. Emirates ID
        id_path = os.path.join(applicant_folder, 'emirates_id.png')
        create_mock_emirates_id(name, id_path)

        # 2. Resume (now in multiple formats)
        resume_content = create_resume_content(name, employment_years)
        create_docx_resume(resume_content, os.path.join(applicant_folder, 'resume.docx'))
        create_pdf_resume(resume_content, os.path.join(applicant_folder, 'resume.pdf'))

        # 3. Bank Statement (as TXT)
        bank_stmt_path = os.path.join(applicant_folder, 'bank_statement.txt')
        with open(bank_stmt_path, 'w') as f:
            f.write(f"Bank Statement for {name}\n")
            f.write(f"Average Monthly Income: AED {income + np.random.randint(-500, 500)}\n")

        # 4. Assets/Liabilities (as Excel)
        assets_path = os.path.join(applicant_folder, 'assets.xlsx')
        assets_df = pd.DataFrame({
            'Asset': ['Savings Account', 'Car'], 'Value': [np.random.randint(1000, 50000), np.random.randint(0, 80000)],
            'Liability': ['Personal Loan', 'Credit Card Debt'], 'Amount': [np.random.randint(0, 20000), np.random.randint(0, 10000)]
        })
        assets_df.to_excel(assets_path, index=False)
        
        data.append({
            'applicant_id': applicant_id,
            'name': name, 'age': np.random.randint(22, 60),
            'monthly_income': income, 'family_size': family_size,
            'employment_years': employment_years,
            'address_form': fake.address().replace('\n', ', '),
            'eligibility_status': eligibility
        })

    df = pd.DataFrame(data)
    df.to_csv('data/applications.csv', index=False)
    print(f"Generated {num_applicants} synthetic applicant profiles in 'data/applicants' and 'data/applications.csv'")

if __name__ == "__main__":
    create_applicant_data(50)